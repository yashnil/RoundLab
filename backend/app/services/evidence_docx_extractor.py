"""DOCX text extraction using python-docx.

SAFETY INVARIANTS
- Extracted text is exact source output; never synthesized.
- Paragraph order is always preserved.
- Heading styles are preserved as section headings.
- Bold/underline/italic runs are tracked in parser_metadata but text is not altered.
- Malformed DOCX returns a gracefully-failed ExtractedDocument.

python-docx is in requirements.txt.
"""

from __future__ import annotations

import io
import logging
from datetime import datetime, timezone

logger = logging.getLogger(__name__)

_HEADING_STYLES = frozenset({
    "heading 1", "heading 2", "heading 3",
    "heading 4", "heading 5", "heading 6",
    "title", "subtitle",
})


def extract_docx(
    source: bytes | str,
    *,
    source_url: str = "",
    retrieval_timestamp: str = "",
) -> "ExtractedDocument":
    """Extract text from a DOCX file.

    Args:
        source: raw DOCX bytes or a local file path.
        source_url: original URL for provenance.
        retrieval_timestamp: ISO-8601 UTC timestamp.

    Returns an ExtractedDocument.  Headings become `section.heading` on the
    following paragraph sections.  Bold/italic/underline metadata is stored
    in `parser_metadata` for each section but text is not modified.
    """
    from app.services.evidence_extracted_document import (
        DocumentSection,
        ExtractedDocument,
        finalize_document,
    )

    ts = retrieval_timestamp or datetime.now(timezone.utc).isoformat()

    doc = ExtractedDocument(
        source_url=source_url,
        document_type="docx",
        extraction_method="python_docx",
        extraction_version="",
        retrieval_timestamp=ts,
        http_content_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
    )

    try:
        from docx import Document  # type: ignore
        import docx  # type: ignore
        doc.extraction_version = f"python-docx-{docx.__version__}"
    except ImportError:
        doc.extraction_warnings.append("python-docx not available; DOCX extraction skipped.")
        doc.source_text_type = "metadata_only"
        doc.is_metadata_only = True
        doc.extraction_confidence = 0.0
        return doc

    try:
        if isinstance(source, bytes):
            docx_doc = Document(io.BytesIO(source))
        else:
            docx_doc = Document(str(source))
    except Exception as exc:
        doc.extraction_warnings.append(f"Could not open DOCX: {exc}")
        doc.source_text_type = "metadata_only"
        doc.is_metadata_only = True
        doc.extraction_confidence = 0.0
        return doc

    try:
        # Extract core properties when available
        try:
            props = docx_doc.core_properties
            if props.title:
                doc.title = props.title.strip()
                doc.metadata_provenance["title"] = "docx_properties"
            if props.author:
                doc.author = props.author.strip()
                doc.metadata_provenance["author"] = "docx_properties"
            if props.created:
                doc.publication_date = props.created.strftime("%Y-%m-%d")
                doc.metadata_provenance["publication_date"] = "docx_properties"
        except Exception:
            pass  # core_properties may not always be accessible

        all_text_parts: list[str] = []
        sections: list[DocumentSection] = []
        char_offset = 0
        section_idx = 0
        para_idx = 0
        current_heading = ""

        for para in docx_doc.paragraphs:
            para_text = para.text.strip()
            if not para_text:
                continue

            style_name = (para.style.name or "").lower() if para.style else ""
            is_heading = style_name in _HEADING_STYLES

            if is_heading:
                current_heading = para_text
                # Headings are NOT added as separate sections; they label the next
                # content section.
                continue

            # Collect run-level formatting metadata
            run_meta: list[dict] = []
            for run in para.runs:
                if not run.text:
                    continue
                rmd: dict = {}
                try:
                    if run.bold:
                        rmd["bold"] = True
                    if run.italic:
                        rmd["italic"] = True
                    if run.underline:
                        rmd["underline"] = True
                except Exception:
                    pass
                if rmd:
                    rmd["text"] = run.text
                    run_meta.append(rmd)

            start = char_offset
            all_text_parts.append(para_text)
            char_offset += len(para_text)
            all_text_parts.append("\n\n")
            char_offset += 2

            section = DocumentSection(
                text=para_text,
                start_char=start,
                end_char=start + len(para_text),
                heading=current_heading,
                paragraph_index=para_idx,
                section_index=section_idx,
                parser_metadata={"runs": run_meta} if run_meta else {},
            )
            sections.append(section)
            section_idx += 1
            para_idx += 1
            current_heading = ""  # heading consumed by first following para

        raw_text = "".join(all_text_parts).strip()
        if not raw_text:
            doc.extraction_warnings.append("DOCX extraction produced no usable text.")
            doc.source_text_type = "metadata_only"
            doc.is_metadata_only = True
            doc.extraction_confidence = 0.0
            return finalize_document(doc)

        doc.raw_text = raw_text
        doc.sections = sections
        doc.extraction_confidence = 0.95
        doc.source_text_type = "full_text"
        return finalize_document(doc)

    except Exception as exc:
        logger.warning("DOCX extraction error for %s: %s", source_url, exc)
        doc.extraction_warnings.append(f"Unexpected extraction error: {exc}")
        doc.source_text_type = "metadata_only"
        doc.is_metadata_only = True
        doc.extraction_confidence = 0.0
        return finalize_document(doc)
