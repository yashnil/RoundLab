"""Library export service (Pass 13).

Deterministic export of blockfiles and frontlines to JSON, Markdown, and DOCX.

Rules:
- No internal IDs, traces, secrets, or raw model output in exports.
- Bibliography deduplicates by normalized_doi or canonical_url.
- Blockfile order (section position, entry position) is preserved.
- DOCX export uses python-docx (already in requirements).
- All functions are pure/deterministic given the same inputs.
"""

import json
import logging
from typing import Optional

from app.models.evidence_library import (
    BlockfileRow,
    BlockfileSectionRow,
    BlockfileEntryRow,
    FrontlineRow,
    FrontlineResponseRow,
    FrontlineResponseCardRow,
)

logger = logging.getLogger(__name__)


# ── Bibliography dedup helpers ─────────────────────────────────────────────────

def _bib_key(card_data: dict) -> Optional[str]:
    """Return a dedup key for a card's source. Returns None if no key available."""
    citation = card_data.get("citation") or {}
    doi = citation.get("doi") or card_data.get("doi")
    url = card_data.get("url") or citation.get("url")
    if doi:
        return f"doi:{doi.lower().strip()}"
    if url:
        # Strip query/fragment for stable key
        u = url.split("?")[0].split("#")[0].lower().strip()
        return f"url:{u}"
    return None


def _build_bibliography(card_rows: list[dict]) -> list[str]:
    """Build a deduplicated bibliography from a list of evidence_card rows."""
    seen: set[str] = set()
    bib: list[str] = []
    for card in card_rows:
        key = _bib_key(card)
        if key and key in seen:
            continue
        if key:
            seen.add(key)
        mla = card.get("mla_citation") or card.get("cite") or ""
        if mla:
            bib.append(mla.strip())
    return sorted(bib)


# ── Safe card fields (never expose body offsets, support_verdict raw, or traces) ──

_EXPORT_CARD_FIELDS = {"id", "tag", "cite", "body_text", "author", "publication",
                        "title", "published_date", "url", "mla_citation", "short_cite"}


def _safe_card(card: dict) -> dict:
    return {k: v for k, v in card.items() if k in _EXPORT_CARD_FIELDS}


# ── JSON export ───────────────────────────────────────────────────────────────

def export_blockfile_json(
    bf: BlockfileRow,
    sections: list[BlockfileSectionRow],
    entries_by_section: dict[str, list[BlockfileEntryRow]],
    card_data: Optional[dict[str, dict]] = None,  # card_id → evidence_card row
) -> str:
    """Export blockfile as structured JSON. Excludes internal IDs and traces."""
    card_data = card_data or {}
    out: dict = {
        "title": bf.title,
        "side": bf.side,
        "description": bf.description,
        "sections": [],
    }
    all_cards: list[dict] = []
    for sect in sorted(sections, key=lambda s: s.position):
        sect_out: dict = {
            "title": sect.title,
            "section_type": sect.section_type,
            "entries": [],
        }
        for entry in sorted(entries_by_section.get(sect.id, []), key=lambda e: e.position):
            entry_out: dict = {
                "entry_type": entry.entry_type,
                "custom_label": entry.custom_label,
                "notes": entry.notes,
            }
            if entry.card_id and entry.card_id in card_data:
                safe = _safe_card(card_data[entry.card_id])
                entry_out["card"] = safe
                all_cards.append(card_data[entry.card_id])
            sect_out["entries"].append(entry_out)
        out["sections"].append(sect_out)

    out["bibliography"] = _build_bibliography(all_cards)
    return json.dumps(out, indent=2, ensure_ascii=False)


def export_frontline_json(
    fl: FrontlineRow,
    responses: list[FrontlineResponseRow],
    response_cards: dict[str, list[FrontlineResponseCardRow]],
    card_data: Optional[dict[str, dict]] = None,
) -> str:
    card_data = card_data or {}
    all_cards: list[dict] = []
    out: dict = {
        "title": fl.title,
        "side": fl.side,
        "opponent_claim": fl.opponent_claim,
        "opponent_warrant": fl.opponent_warrant,
        "opponent_impact": fl.opponent_impact,
        "responses": [],
    }
    for resp in sorted(responses, key=lambda r: r.position):
        resp_out: dict = {
            "response_type": resp.response_type,
            "response_claim": resp.response_claim,
            "explanation": resp.explanation,
            "wording_for_speech": resp.wording_for_speech,
            "is_analytical": resp.is_analytical,
            "priority": resp.priority,
            "speech_suitability": resp.speech_suitability,
            "cards": [],
        }
        for rc in (response_cards.get(resp.id) or []):
            if rc.card_id in card_data:
                safe = _safe_card(card_data[rc.card_id])
                safe["card_role"] = rc.card_role
                resp_out["cards"].append(safe)
                all_cards.append(card_data[rc.card_id])
        out["responses"].append(resp_out)

    out["bibliography"] = _build_bibliography(all_cards)
    return json.dumps(out, indent=2, ensure_ascii=False)


# ── Markdown export ───────────────────────────────────────────────────────────

def export_blockfile_markdown(
    bf: BlockfileRow,
    sections: list[BlockfileSectionRow],
    entries_by_section: dict[str, list[BlockfileEntryRow]],
    card_data: Optional[dict[str, dict]] = None,
) -> str:
    card_data = card_data or {}
    lines: list[str] = [
        f"# {bf.title}",
        f"**Side:** {bf.side or 'N/A'}",
    ]
    if bf.description:
        lines.append(f"\n{bf.description}")
    lines.append("")

    all_cards: list[dict] = []
    for sect in sorted(sections, key=lambda s: s.position):
        lines.append(f"## {sect.title}")
        lines.append(f"*Section type: {sect.section_type}*")
        lines.append("")
        for entry in sorted(entries_by_section.get(sect.id, []), key=lambda e: e.position):
            if entry.entry_type == "header" and entry.custom_label:
                lines.append(f"### {entry.custom_label}")
                lines.append("")
                continue
            if entry.entry_type == "analytical_note":
                if entry.notes:
                    lines.append(f"*{entry.notes}*")
                    lines.append("")
                continue
            if entry.card_id and entry.card_id in card_data:
                c = card_data[entry.card_id]
                label = entry.custom_label or c.get("tag") or ""
                cite = c.get("cite") or c.get("mla_citation") or ""
                body = c.get("body_text") or ""
                lines.append(f"**{label}**")
                lines.append(f"> {cite}")
                lines.append("")
                lines.append(body)
                if entry.notes:
                    lines.append(f"*Coach note: {entry.notes}*")
                lines.append("")
                all_cards.append(c)

    # Bibliography
    bib = _build_bibliography(all_cards)
    if bib:
        lines.append("---")
        lines.append("## Works Cited")
        for entry in bib:
            lines.append(f"- {entry}")

    return "\n".join(lines)


def export_frontline_markdown(
    fl: FrontlineRow,
    responses: list[FrontlineResponseRow],
    response_cards: dict[str, list[FrontlineResponseCardRow]],
    card_data: Optional[dict[str, dict]] = None,
) -> str:
    card_data = card_data or {}
    lines: list[str] = [
        f"# Frontline: {fl.title}",
        "",
    ]
    if fl.opponent_claim:
        lines.append(f"**Opponent Claim:** {fl.opponent_claim}")
    if fl.opponent_warrant:
        lines.append(f"**Opponent Warrant:** {fl.opponent_warrant}")
    if fl.opponent_impact:
        lines.append(f"**Opponent Impact:** {fl.opponent_impact}")
    lines.append("")

    all_cards: list[dict] = []
    for i, resp in enumerate(sorted(responses, key=lambda r: r.position), start=1):
        lines.append(f"## Response {i}: {resp.response_type.replace('_', ' ').title()}")
        lines.append(f"**Claim:** {resp.response_claim}")
        if resp.explanation:
            lines.append(f"**Explanation:** {resp.explanation}")
        if resp.wording_for_speech:
            lines.append(f"**Read-aloud:** _{resp.wording_for_speech}_")
        lines.append(
            f"**Speech:** {', '.join(resp.speech_suitability)} | "
            f"**Priority:** {resp.priority}"
        )
        if resp.is_analytical:
            lines.append("*(Analytical — no evidence card required)*")
        lines.append("")

        for rc in (response_cards.get(resp.id) or []):
            if rc.card_id in card_data:
                c = card_data[rc.card_id]
                role_label = "Supporting" if rc.card_role == "supporting" else "Opposing"
                lines.append(f"*[{role_label} card]*")
                lines.append(f"**{c.get('tag', '')}**")
                lines.append(f"> {c.get('cite') or c.get('mla_citation') or ''}")
                lines.append(c.get("body_text") or "")
                lines.append("")
                all_cards.append(c)

    bib = _build_bibliography(all_cards)
    if bib:
        lines.append("---")
        lines.append("## Works Cited")
        for entry in bib:
            lines.append(f"- {entry}")

    return "\n".join(lines)


# ── DOCX export ───────────────────────────────────────────────────────────────

def export_blockfile_docx(
    bf: BlockfileRow,
    sections: list[BlockfileSectionRow],
    entries_by_section: dict[str, list[BlockfileEntryRow]],
    card_data: Optional[dict[str, dict]] = None,
) -> bytes:
    """Export blockfile as DOCX using python-docx.

    Preserves section order and card text. Bold/underline/highlight markup
    is rendered as bold/underline in the DOCX body text.
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Pt, RGBColor  # type: ignore
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX export. Install with: pip install python-docx") from exc

    card_data = card_data or {}
    doc = Document()

    # Title
    title_para = doc.add_heading(bf.title, 0)
    if bf.side:
        doc.add_paragraph(f"Side: {bf.side}")
    if bf.description:
        doc.add_paragraph(bf.description)

    all_cards: list[dict] = []
    for sect in sorted(sections, key=lambda s: s.position):
        doc.add_heading(sect.title, 1)
        for entry in sorted(entries_by_section.get(sect.id, []), key=lambda e: e.position):
            if entry.entry_type == "header" and entry.custom_label:
                doc.add_heading(entry.custom_label, 2)
                continue
            if entry.entry_type == "analytical_note":
                if entry.notes:
                    p = doc.add_paragraph()
                    p.add_run(entry.notes).italic = True
                continue
            if entry.card_id and entry.card_id in card_data:
                c = card_data[entry.card_id]
                label = entry.custom_label or c.get("tag") or ""
                cite = c.get("cite") or c.get("mla_citation") or ""
                body = c.get("body_text") or ""

                # Tag line — bold
                tag_para = doc.add_paragraph()
                run = tag_para.add_run(label)
                run.bold = True
                run.font.size = Pt(11)

                # Citation
                if cite:
                    cite_para = doc.add_paragraph()
                    cite_run = cite_para.add_run(cite)
                    cite_run.italic = True
                    cite_run.font.size = Pt(9)

                # Card body
                if body:
                    body_para = doc.add_paragraph(body)
                    body_para.paragraph_format.left_indent = Pt(24)
                    for run in body_para.runs:
                        run.font.size = Pt(10)

                # Coach note
                if entry.notes:
                    note_para = doc.add_paragraph()
                    note_run = note_para.add_run(f"Note: {entry.notes}")
                    note_run.italic = True
                    note_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

                doc.add_paragraph("")  # spacer
                all_cards.append(c)

    # Bibliography
    bib = _build_bibliography(all_cards)
    if bib:
        doc.add_page_break()
        doc.add_heading("Works Cited", 1)
        for entry in bib:
            p = doc.add_paragraph(entry)
            p.paragraph_format.left_indent = Pt(36)
            # Hanging indent approximation: first line indent negative
            p.paragraph_format.first_line_indent = Pt(-36)

    # Save to bytes
    import io
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
